from docx import Document
import requests
from deep_translator import GoogleTranslator


def generate_tasks(students_name, students_class, subject, amount_of_tasks, list_of_themes, prompt):
    url = "http://localhost:1337/v1/chat/completions"
    body = {
        "model": "gpt-3.5-turbo-16k",
        "stream": False,
        'messages': [{'role': 'system',
                      'content': f'You are a good {subject} teacher in school in {students_class} grade. {prompt} '},
                     {'role': 'user',
                      'content': f'Generate{amount_of_tasks} {subject} tasks for {students_name} in {students_class} grade'
                                 f'who does not understand following themes: {list_of_themes}. Also generate tasks on '
                                 f'close themes to revise material, but dont forget that students are in '
                                 f'{students_class} grade '}]}
    json_response = requests.post(url, json=body).json().get('choices', [])
    return json_response


def return_generated_tasks_docx(dataframe, class_level, additional_prompt):
    doc = Document()
    file_name = "Готовые_задания.docx"
    for i in range(dataframe.shape[0]):
        generated_content = generate_tasks(students_name=dataframe["ФИО Ученика"][i], students_class=class_level,
                                           subject=dataframe["Предмет"][i],
                                           amount_of_tasks=dataframe["Количество заданий"][i],
                                           list_of_themes=dataframe["Темы"][i], prompt=additional_prompt)
        for choice in generated_content:
            language = dataframe["Язык"][i]
            if language:
                doc.add_paragraph(GoogleTranslator(source='auto', target=dataframe["Язык"][i]).translate(
                    choice.get('message', {}).get('content', '')))
            else:
                doc.add_paragraph(GoogleTranslator(source='auto', target='ru').translate(
                    choice.get('message', {}).get('content', '')))
    doc.save(file_name)
    return file_name
